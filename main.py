from flair.models import SequenceTagger
from flair.data import Sentence
import docx
from docx.oxml.ns import qn
from translate import Translator
import openai
import streamlit as st
from streamlit_option_menu import option_menu
import warnings
warnings.filterwarnings("ignore")


st.set_page_config(page_title="DocTranslator", page_icon="",layout="wide")
st.markdown(
    """
    <style>
    .main {
        padding: 0rem 0rem;
    }
    </style>
    """,
    unsafe_allow_html=True
)
st.title("Docx Translator")


tagger = SequenceTagger.load('ner')


openai.api_key  = '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'
def name_extractor(doc):
    name_list = []
    for para in doc.paragraphs:
        sentence = Sentence(para.text)
        tagger.predict(sentence)
        for entity in sentence.get_spans('ner'):
            if entity.tag == 'PER' and entity.text not in name_list:
                name_list.append(entity.text)
    return name_list

def remove_short_names(names):
    filtered_names = []
    for name in names:
        a = names.copy()
        a.remove(name)
        b = ",".join(a)
        if name not in b:
          filtered_names.append(name)
    return filtered_names


def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
    model=model,
    messages=messages,
    temperature=0,  # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]

def copy_run_properties(source_run, target_run):
    # Copy font properties
    target_run.font.size = source_run.font.size
    target_run.font.name = source_run.font.name
    target_run.font.color.rgb = source_run.font.color.rgb

    # Copy other properties as needed
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    # Add more properties as needed


def translate_docx(doc, output_path, target_language, new_names):

    translator = Translator(to_lang=target_language)
    translated_doc = docx.Document()
    current_list_style = None
    current_list_level = 0
    for para in doc.paragraphs:
        if para.style.name.startswith('List'):
            list_style = para.style.name.split(' ')[0]
            list_level = int(para.style.name.split(' ')[1])

            # Create a new list item in the translated document
            new_list_item = translated_doc.add_paragraph()
            new_list_item.style = list_style

            # Indent the list item based on the level
            indent_size = list_level * 0.5  # Adjust the indentation as needed

            # Create a new run for the list item's text
            new_run = new_list_item.add_run()

            # Apply indentation to the run's XML properties
            new_run._element.get_or_add_tcPr().append(
              docx.oxml.ns.qn('w:ind'),
              {
                'left': str(int(indent_size * 720)),  # Convert inches to twips
                'hanging': '0'
              }
            )

            # Check if the list style or level has changed
            if list_style != current_list_style or list_level != current_list_level:
                new_list = new_list_item.add_run().add_field('listNum', f'{list_level + 1}')
                current_list_style = list_style
                current_list_level = list_level
            else:
                new_list = new_list_item.add_run().add_field('listNum', '\t')
                replaced_text = para.text
                for original_name, replacement_name in new_names.items():
                      replaced_text = replaced_text.replace(original_name, replacement_name)

                translated_text = translator.translate(replaced_text)
                new_list_item.add_run(translated_text)

        else:
            # Handle non-list paragraphs
            new_para = translated_doc.add_paragraph()
            new_para.alignment = para.alignment

            # Copy run-level formatting
            for run in para.runs:
                new_run = new_para.add_run()
                copy_run_properties(run, new_run)
                replaced_text = run.text
                for original_name, replacement_name in new_names.items():
                    replaced_text = replaced_text.replace(original_name, replacement_name)
                new_run.text = translator.translate(replaced_text)

    # Save the translated content to a new DOCX file
    return translated_doc

selected = option_menu(
        menu_title=None,  # required
        options=["Home", "About"],  # required
        icons=["house", "envelope"],  # optional
        menu_icon="cast",  # optional
        default_index=0,  # optional
styles = {"nav-link": {"--hover-color": "grey"}},
        orientation="horizontal",
    )
# Homepage content upon clicking Home option in the tab
if selected == "Home":

    input_path = st.file_uploader(label="Upload Docx Format File",type=['docx'])
    language = {'English': 'en', 'Spanish': 'es', 'French': 'fr', 'German': 'de', 'Italian': 'it',
                'Chinese': 'zh', 'Japanese': 'ja', 'Russian': 'ru', 'Portuguese': 'pt', 'Dutch': 'nl',
                'Arabic': 'ar', 'Hindi': 'hi'}
    doc = docx.Document(input_path)
    c1,c2,c3 = st.columns(3)
    with c1:
        lang = st.selectbox(label="Choose Language",options=language.keys())
        target_language = language[lang]
    with c2:
        region = {
        "India": "Indian",
        "United States": "American",
        "China": "Chinese",
        "Brazil": "Brazilian",
        "Russia": "Russian",
        "Indonesia": "Indonesian",
        "Pakistan": "Pakistani",
        "Nigeria": "Nigerian",
        "Bangladesh": "Bangladeshi",
        "Japan": "Japanese",
        "Mexico": "Mexican",
        "Philippines": "Filipino",
        "Egypt": "Egyptian",
        "Vietnam": "Vietnamese",
        "Ethiopia": "Ethiopian",
    }
        reg = st.selectbox(label="Choose Language", options=region.keys())
        reg_select = region[reg]

    with c3:
        upload = st.button("Upload")
    if upload:
        name_list = name_extractor(doc)
        filtered_names = remove_short_names(name_list)
        new_names = dict()
        for name in filtered_names:
            prompt = f"""give a alternative {reg_select} name for {name}, 
            give only one name without any other word in English letters only"""
            response = get_completion(prompt)
            for m,n in zip(name.split(),response.split()):
                new_names[m] = n
        output_docx_path = "Translated Document.docx"

        final_doc = translate_docx(doc, output_docx_path, target_language,new_names)
        final_doc.save(output_docx_path)  # Save the translated document to a file

        with open(output_docx_path, "rb") as f:
            st.download_button(label='Download', data=f, file_name=output_docx_path)

elif selected == "About":
    st.markdown('__<p style="text-align:left; font-size: 25px; color: #FAA026">Summary of Docx Translator Project</P>__',
                unsafe_allow_html=True)
    st.write("This Docx Translator project focused on the users to translate the docx from one language to another along with name localization.")
    st.markdown('__<p style="text-align:left; font-size: 20px; color: #FAA026">Applications and Packages Used:</P>__',
                    unsafe_allow_html=True)
    st.write("  * Python")
    st.write("  * SequenceTagger Sentence from flair.models")
    st.write("  * Streamlit and and streamlit_option_menu")
    st.write("  * Docx")
    st.write("  * Translator from translate ")
    st.write("  * Openai")
    st.write("  * Github")
    st.markdown('__<p style="text-align:left; font-size: 20px; color: #FAA026">For feedback/suggestion, connect with me on</P>__',
                unsafe_allow_html=True)
    st.subheader("LinkedIn")
    st.write("https://www.linkedin.com/in/selvamani-a-795580266/")
    st.subheader("Email ID")
    st.write("selvamani.ind@gmail.com")
    st.subheader("Github")
    st.write("https://github.com/selvamani1992")
    #st.balloons()